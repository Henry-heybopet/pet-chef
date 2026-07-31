from pathlib import Path
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUT_DOCX = ROOT / "Heybo_Pet_中文版用户操作说明书.docx"

GREEN = "157A52"
TEAL = "178A83"
ORANGE = "B44F25"
INK = "27312C"
GRAY = "657069"
PALE_GREEN = "EAF6F0"
PALE_ORANGE = "FFF3E7"
PALE_BLUE = "EAF7F8"


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, east_asia="Arial Unicode MS", ascii_font="Arial Unicode MS"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_callout(doc, title, body, tone="info"):
    colors = {
        "info": (PALE_BLUE, TEAL),
        "warn": (PALE_ORANGE, ORANGE),
        "tip": (PALE_GREEN, GREEN),
    }
    fill, accent = colors[tone]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.right_indent = Cm(0.2)
    set_cell_shading(p, fill)
    r = p.add_run(f"{title}\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    set_run_font(r)
    r = p.add_run(body)
    r.font.color.rgb = RGBColor.from_string(INK)
    set_run_font(r)
    return p


def add_step(doc, number, title, body):
    p = doc.add_paragraph(style="Manual Step")
    r = p.add_run(f"{number}. {title}：")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    set_run_font(r)
    r = p.add_run(body)
    set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Manual Bullet")
    p.paragraph_format.left_indent = Cm(0.55 + level * 0.55)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run("• " + text)
    set_run_font(r)
    return p


def add_figure(doc, filename, caption, width=3.25):
    path = SHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return cap


def prepare_long_screenshot_parts():
    """Split long phone screenshots so every part remains readable in Word/PDF."""
    split_specs = {
        "app-recipe-list-stage-actual.jpg": [0, 1750, 4385, 6620, 8380],
        "app-recipe-detail-actual.jpg": [0, 2170, 4400, 5895],
        "app-fresh-check-result-actual.jpg": [0, 3360, 4780, 7240, 9807],
    }
    for filename, boundaries in split_specs.items():
        source = SHOTS / filename
        with Image.open(source) as image:
            for index, (top, bottom) in enumerate(zip(boundaries, boundaries[1:]), start=1):
                part = image.crop((0, top, image.width, bottom))
                part.save(
                    SHOTS / f"{source.stem}-part-{index}.jpg",
                    quality=94,
                    subsampling=0,
                )


def chapter(doc, title):
    doc.add_page_break()
    doc.add_heading(title, level=1)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color, before, after in [
        ("Title", 30, GREEN, 0, 8),
        ("Subtitle", 14, GRAY, 0, 8),
        ("Heading 1", 20, GREEN, 0, 14),
        ("Heading 2", 15, TEAL, 12, 6),
        ("Heading 3", 12, ORANGE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    manual_step = styles.add_style("Manual Step", WD_STYLE_TYPE.PARAGRAPH)
    manual_step.base_style = styles["Normal"]
    manual_step.paragraph_format.left_indent = Cm(0.15)
    manual_step.paragraph_format.space_after = Pt(6)
    manual_step.paragraph_format.keep_together = True

    manual_bullet = styles.add_style("Manual Bullet", WD_STYLE_TYPE.PARAGRAPH)
    manual_bullet.base_style = styles["Normal"]
    manual_bullet.paragraph_format.space_after = Pt(3)

    caption = styles["Caption"]
    caption.name = "Figure Caption"
    caption.font.name = "Arial Unicode MS"
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.font.italic = False
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.space_after = Pt(7)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Heybo Pet｜中文版用户操作说明书")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    set_run_font(r)
    add_page_number(sec.footer.paragraphs[0])
    return doc


def build():
    prepare_long_screenshot_parts()
    doc = setup_document()

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HEYBO PET")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    set_run_font(r)
    t = doc.add_paragraph(style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run("中文版用户操作说明书")
    st = doc.add_paragraph(style="Subtitle")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run("宠物数字档案 · AI 鲜食营养 · 智能鲜食机")
    add_figure(doc, "app-home.jpg", "手机 APP 首页实拍（示例账号名称仅用于界面展示）", width=2.75)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用：Pet Chef Android APP 2.0.0\n文档版本：V1.0｜2026 年 7 月")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    set_run_font(r)

    doc.add_page_break()
    doc.add_heading("使用前必读", level=1)
    add_callout(
        doc,
        "营养与健康提示",
        "APP 的 AI 分析用于辅助理解宠物档案和鲜食方案，不能代替兽医诊断。若宠物处于疾病治疗、术后恢复、妊娠、哺乳，或存在严重过敏、肾脏/胰腺等问题，请先咨询执业兽医或宠物营养专业人士。",
        "warn",
    )
    add_callout(
        doc,
        "设备安全提示",
        "烹饪前请确认锅体、刀组、上盖安装正确，食材未超过容量上限，宠物与儿童远离工作中的机器。不要在机器运行时伸手接触刀组或打开上盖。",
        "warn",
    )
    doc.add_heading("说明书中的示例", level=2)
    add_bullet(doc, "“手册演示”为测试账号；“茜茜”为虚构的示例宠物档案。")
    add_bullet(doc, "所有界面图均来自手机 APP 或 Android APP 模拟器，不使用 Web 页面截图。")
    add_bullet(doc, "不同 Android/iOS 版本的系统权限弹窗文字可能略有差异，以手机实际显示为准。")
    add_bullet(doc, "截图中的 Wi‑Fi 名称、设备名称、宠物数据和食谱仅用于说明操作。")

    doc.add_heading("目录", level=1)
    for item in [
        "第 1 章　手机号短信验证码注册与登录",
        "第 2 章　建立与完善宠物数字档案",
        "第 3 章　保存宠物档案并使用 AI 查询",
        "第 4 章　阅读 AI 营养需求分析并选择食谱",
        "第 5 章　筛选食谱并快速进入一键烹饪",
        "第 6 章　绑定鲜食机：权限、2.4G Wi‑Fi 与配网",
        "第 7 章　鲜食评估：粘贴识别或手动输入食谱",
        "第 8 章　查看 AI 食谱推荐，理解宠物与食谱的关系",
        "第 9 章　一键烹饪与喂食反馈",
        "附录　常见问题与故障排查",
    ]:
        add_bullet(doc, item)

    # Chapter 1
    chapter(doc, "第 1 章　手机号短信验证码注册与登录")
    doc.add_heading("1.1 注册前准备", level=2)
    add_bullet(doc, "准备一个可正常使用的中国大陆 11 位手机号。")
    add_bullet(doc, "准备 1–18 个字符的用户名，可使用中文、英文字母和数字，但不能全部为数字。")
    add_bullet(doc, "阅读《用户协议》和《隐私政策》；只有勾选同意后才能获取验证码并登录。")
    add_figure(doc, "app-login.png", "图 1-1　手机 APP 登录/注册弹窗")
    doc.add_heading("1.2 新手机号注册", level=2)
    add_step(doc, 1, "打开登录窗口", "启动 APP，在首页左上角点击“登录”。若从宠物档案、AI 鲜食评估或烹饪入口进入，APP 也会自动提示先登录。")
    add_step(doc, 2, "输入新手机号", "在手机号栏输入 11 位中国大陆手机号。")
    add_step(doc, 3, "同意协议", "勾选“我已阅读并同意《用户协议》和《隐私政策》”。未注册手机号验证后会自动创建账号。")
    add_step(doc, 4, "获取验证码", "点击“获取验证码”，查收 6 位短信验证码。验证码 5 分钟内有效，60 秒后才可重新发送。")
    add_step(doc, 5, "验证手机号", "输入短信中的 6 位验证码，点击“验证并登录”。系统确认该手机号尚未建立账号后，会进入“设置用户名”页面。")
    add_step(doc, 6, "设置用户名", "输入 1–18 个中文、英文或数字组合；不能含空格、标点或特殊符号，也不能全部是数字。")
    add_step(doc, 7, "完成登录", "点击“完成登录”。首页左上角由“登录”变为用户名，表示注册与登录成功。")
    add_callout(doc, "验证码说明", "验证码连续输入错误 5 次后会失效，需要重新获取。不要把收到的验证码告诉任何人。", "tip")
    doc.add_heading("1.3 已有账号登录", level=2)
    add_step(doc, 1, "输入手机号", "输入账号已经绑定的 11 位中国大陆手机号。用户名不再用于登录。")
    add_step(doc, 2, "获取并输入验证码", "勾选同意协议，点击“获取验证码”，再输入短信中的 6 位数字。")
    add_step(doc, 3, "验证并登录", "点击“验证并登录”。登录成功后，APP 会加载此账号名下的宠物档案、设备和使用记录；30 天内通常不需要再次使用短信验证码。")
    add_figure(doc, "app-demo-home.png", "图 1-2　登录后的 APP 首页（“手册演示”为测试账号）")
    doc.add_heading("1.4 无法登录时", level=2)
    add_bullet(doc, "确认输入的是账号已绑定的完整 11 位手机号。")
    add_bullet(doc, "确认验证码为短信中最新的 6 位数字，并且未超过 5 分钟。")
    add_bullet(doc, "检查手机网络；切换移动数据/Wi‑Fi 后重试。")
    add_bullet(doc, "验证码连续错误 5 次会失效；请重新获取，不要继续提交旧验证码。")

    # Chapter 2
    chapter(doc, "第 2 章　建立与完善宠物数字档案")
    doc.add_heading("2.1 进入宠物档案", level=2)
    add_step(doc, 1, "进入宠物页", "点击首页“宠物数字档案”，或点击底部“宠物”。")
    add_step(doc, 2, "新增宠物", "在“宠物管理主页”点击“+ 添加我的宠物”。一个账号可为多只宠物分别建立档案。")
    add_figure(doc, "app-pet-list.png", "图 2-1　宠物管理主页；点击上方“添加我的宠物”新建档案")

    doc.add_heading("2.2 上传宠物照片", level=2)
    add_step(doc, 1, "点击上传区", "在基础档案左上方点击圆形“上传”区域。")
    add_step(doc, 2, "选择照片来源", "根据手机提示选择相册中的照片；若首次使用相册或相机，请允许所需权限。")
    add_step(doc, 3, "选择清晰正面照", "优先使用光线充足、主体完整、没有过多遮挡的照片。")
    add_step(doc, 4, "检查预览", "返回档案页后确认圆形头像已更新；不满意可再次点击更换。")
    add_callout(doc, "隐私建议", "不要上传包含家庭住址、证件、车牌或其他个人信息的照片。", "info")
    add_figure(doc, "app-pet-profile-basic-actual.jpg", "图 2-2　基础档案：上传照片，并填写姓名、性别、生日、品种、体型、活动水平、体重、BCS 与喂养目标", width=2.65)

    doc.add_heading("2.3 逐项填写基础档案", level=2)
    add_step(doc, 1, "宠物姓名", "填写日常使用的名字，方便在 AI 分析、食谱和喂食记录中识别。")
    add_step(doc, 2, "宠物性别", "选择公、母或隐私选项。性别会影响特殊生理状态可选项。")
    add_step(doc, 3, "出生日期", "尽量填写准确日期；不确定时可根据领养记录或兽医估计年龄填写。")
    add_step(doc, 4, "选择品种", "点击“搜索品种”，输入品种关键词，从下拉候选中点选。若没有匹配项，可选择其他/自定义并填写实际品种。")
    add_step(doc, 5, "选择体型", "按宠物成年体型选择迷你、小型、中型、大型或巨型；不要只凭当前体重判断幼犬成年体型。")
    add_step(doc, 6, "活动水平", "低活动适合久坐或运动很少；中度适合普通家庭犬；高活动/工作犬适合规律高强度运动。")
    add_step(doc, 7, "喂养环境", "根据主要生活环境选择室内、室外或混合。")
    add_step(doc, 8, "当前与目标体重", "使用近期称重数据，单位为 kg。目标体重应参考体况评分和兽医建议。")
    add_step(doc, 9, "BCS 体况", "按 APP 引导评估肋骨触感、腰线和腹线；BCS 5/9 通常代表理想体况。")
    add_step(doc, 10, "喂养目标", "选择维持体重、减重、增肌、美毛、肠胃照护或术后恢复等最主要目标。")
    add_callout(doc, "品种选择要点", "请从候选列表中真正点选品种，不能只停留在输入文字状态。混种犬可选择最接近的体型/主品种，并在健康信息中补充说明。", "tip")

    doc.add_heading("2.4 填写健康档案", level=2)
    add_step(doc, 1, "切换到健康档案", "点击顶部“② 健康档案”。")
    add_step(doc, 2, "绝育/特殊时期", "选择已绝育、妊娠期、哺乳期、术后休养、病后恢复或无特殊情况；公犬不能选择妊娠或哺乳。")
    add_step(doc, 3, "过敏原", "填写已确认或高度怀疑的食物，如鸡肉、牛肉、乳制品等。多个项目请分开记录。")
    add_step(doc, 4, "过敏表现", "记录红斑、瘙痒、软便、呕吐等表现，并选择轻微、中等或严重。")
    add_step(doc, 5, "健康历史", "按实际情况多选肥胖、心脏、肾脏、肝脏、肠胃、泌尿、皮肤、关节等问题。")
    add_figure(doc, "app-pet-profile-health-actual.jpg", "图 2-3　健康档案：特殊时期与绝育状态、过敏原、过敏表现及健康历史", width=3.0)
    add_callout(doc, "准确性优先", "不知道的项目不要猜测。可先留空，完成体检或咨询兽医后再回来修改。错误的过敏或疾病信息可能影响 AI 推荐结果。", "warn")

    # Chapter 3
    chapter(doc, "第 3 章　保存宠物档案并使用 AI 查询")
    doc.add_heading("3.1 保存档案", level=2)
    add_step(doc, 1, "检查必填项", "确认姓名、性别、生日、品种、体型、活动水平和体重已填写。")
    add_step(doc, 2, "返回并检查健康档案", "特别检查过敏原和健康历史，避免遗漏。")
    add_step(doc, 3, "点击保存", "保存成功后返回宠物管理主页，并显示宠物卡片。")
    add_step(doc, 4, "核对卡片", "确认姓名、年龄、品种、体重、BCS、目标和过敏原与实际一致。")
    add_callout(doc, "为什么必须先保存", "AI 查询使用已保存的 pet_id 和档案数据。尚未保存的草稿不能作为稳定的 AI 分析对象。", "info")
    add_figure(doc, "app-pet-list.png", "图 3-1　已保存档案会显示宠物卡片；点击卡片进入 AI 查询")

    doc.add_heading("3.2 使用 AI 查询宠物档案", level=2)
    add_step(doc, 1, "点击宠物卡片", "在“我的宠物”列表中点击要分析的宠物，不要点击左侧“修改”。")
    add_step(doc, 2, "等待分析", "APP 会读取所选宠物的年龄、品种、体重、BCS、活动水平、目标、过敏与健康信息。")
    add_step(doc, 3, "查看结果", "页面显示每日总量、每日餐次、每餐份量、能量、核心营养需求、AI 建议和注意事项。")
    add_step(doc, 4, "需要更新时先改档案", "若体重、健康状况或过敏信息改变，先回到宠物列表点击“修改”，保存后重新分析。")
    add_figure(doc, "app-ai-analysis.png", "图 3-2　AI 营养分析首页；数据与当前选择的宠物档案对应")

    # Chapter 4
    chapter(doc, "第 4 章　阅读 AI 营养需求分析并选择合适食谱")
    doc.add_heading("4.1 先确认“分析对象”", level=2)
    add_bullet(doc, "顶部应显示正确品种、当前体重、年龄阶段和活动水平。")
    add_bullet(doc, "若分析对象不对，返回宠物管理主页重新选择宠物。")
    add_bullet(doc, "多宠物家庭不能把一只宠物的分析直接用于另一只宠物。")
    doc.add_heading("4.2 逐区阅读分析", level=2)
    add_step(doc, 1, "每日总量", "是当前档案条件下的每日鲜食参考总克数，不是单餐克数。")
    add_step(doc, 2, "每日餐次与每餐份量", "例如每日 610g、2 餐，则每餐约 305g。实际喂食需结合零食、其他食物和体重变化调整。")
    add_step(doc, 3, "每日能量", "以 kcal 表示，是用于理解能量需求的参考值。")
    add_step(doc, 4, "食谱能量密度", "单位 kcal/g；密度更高并不一定更适合，需结合每日总量和宠物目标。")
    add_step(doc, 5, "核心营养需求", "显示优质蛋白、脂肪酸、纤维等重点，是筛选食谱的方向。")
    add_step(doc, 6, "AI 营养师建议", "阅读总体建议和调整依据。")
    add_step(doc, 7, "注意事项", "橙色警示优先级高，尤其是过敏原、健康限制或特殊生理期提示。")
    add_figure(doc, "app-ai-analysis.png", "图 4-1　每日总量、餐次、能量、核心需求与注意事项")

    doc.add_heading("4.3 理解 A+B+C 组合", level=2)
    add_bullet(doc, "A 鲜食基础包（必选）：主要食材组成和基础配方。")
    add_bullet(doc, "B 全价营养包（必选）：用于补足基础食材难以稳定覆盖的维生素、矿物质等。")
    add_bullet(doc, "C 功能支持包（可选）：围绕美毛、关节、肠胃、护肝等目标提供额外支持。")
    add_step(doc, 1, "先检查 A 包", "点击“查看详情”，检查全部食材及比例，排除已知过敏原。")
    add_step(doc, 2, "确认 B 包", "优先保留系统推荐；更换时阅读适用阶段和禁忌。")
    add_step(doc, 3, "按需选择 C 包", "最多选择一个主要功能方向，避免无目的叠加。")
    add_step(doc, 4, "进入制作", "确认组合后点击底部制作按钮。")
    add_figure(doc, "app-ai-recipe-recommendation.png", "图 4-2　AI 推荐的 A 基础包、B 营养包与 C 功能支持包")
    add_callout(doc, "选择原则", "先安全、再适配、后偏好：先排除过敏和健康风险，再匹配年龄/体型/目标，最后考虑宠物口味。", "tip")

    # Chapter 5
    chapter(doc, "第 5 章　筛选不同食谱并快速进入一键烹饪")
    doc.add_heading("5.1 三种筛选方式", level=2)
    add_bullet(doc, "按年龄生命阶段：幼犬、控钙大型幼犬、成犬、老年犬。")
    add_bullet(doc, "按食材种类：鸡肉、牛肉、鱼肉或其他肉类蛋白。")
    add_bullet(doc, "按健康功能：美毛护肤、护肝、低敏单一蛋白、低脂、关节支持等。")
    add_figure(doc, "app-recipe-categories-actual.jpg", "图 5-1　食谱分类中心；可按年龄生命阶段、食材种类和健康功能筛选", width=2.45)
    add_step(doc, 1, "选择一个主要筛选维度", "例如已知鸡肉过敏时，不要从鸡肉蛋白分类开始；可优先选择低敏或其他蛋白。")
    add_step(doc, 2, "进入分类列表", "点击对应分类卡片，APP 显示该分类下的食谱数量。")
    add_step(doc, 3, "展开查看", "点击食谱卡片右侧“展开”，查看食材构成和比例。")
    add_step(doc, 4, "选择食谱", "确认无过敏原、符合阶段与健康目标后，点击右侧“选择”。")
    add_figure(doc, "app-recipe-list-health-actual.jpg", "图 5-2　“美毛护肤”分类食谱；可点击“展开”查看组成，或点击“选择”使用", width=2.4)
    add_figure(doc, "app-recipe-list-stage-actual-part-1.jpg", "图 5-3a　“幼犬通用”分类食谱（上）：展开后可核对食材名称及配方比例", width=3.0)
    add_figure(doc, "app-recipe-list-stage-actual-part-2.jpg", "图 5-3b　“幼犬通用”分类食谱（中上）", width=3.0)
    add_figure(doc, "app-recipe-list-stage-actual-part-3.jpg", "图 5-3c　“幼犬通用”分类食谱（中下）", width=3.0)
    add_figure(doc, "app-recipe-list-stage-actual-part-4.jpg", "图 5-3d　“幼犬通用”分类食谱（下）", width=3.0)

    doc.add_heading("5.2 查看食谱详情并准备制作", level=2)
    add_step(doc, 1, "确认宠物", "制作页顶部显示当前宠物体重、年龄和品种。")
    add_step(doc, 2, "检查适配信息", "查看适配度、推荐标记和适用的鲜食基础包，并确认与当前宠物相符。")
    add_step(doc, 3, "核对配方食材", "逐项检查食材名称、比例和换算克数，确认不存在已知过敏原或禁忌食材。")
    add_step(doc, 4, "阅读营养估算", "查看每日建议能量、食谱能量密度、蛋白食材占体重比例，以及各食材功效说明。")
    add_step(doc, 5, "查看鲜食验证结果", "结合食材安全性、宠物适配性、结构平衡性、营养完整性和长期适宜性判断是否使用。")
    add_step(doc, 6, "进入制作", "确认食谱后返回推荐列表并选择该食谱，再按页面提示选择份数、称量食材、加入清水。")
    add_step(doc, 7, "进入一键烹饪", "所有食材准备完成后点击底部“一键烹饪”。")
    add_figure(doc, "app-recipe-detail-actual-part-1.jpg", "图 5-4a　食谱详情（上）：适配信息、配方比例与每日营养需求估算", width=3.0)
    add_figure(doc, "app-recipe-detail-actual-part-2.jpg", "图 5-4b　食谱详情（中）：各项食材的功效与营养说明", width=3.0)
    add_figure(doc, "app-recipe-detail-actual-part-3.jpg", "图 5-4c　食谱详情（下）：鲜食验证结果", width=3.0)
    add_figure(doc, "app-one-click-cooking.png", "图 5-5　页面底部“一键烹饪”入口")
    add_callout(doc, "称量要求", "以电子秤实测克数为准。不要把百分比直接当作克数，也不要凭目测替代称量。", "warn")

    # Chapter 6
    chapter(doc, "第 6 章　绑定鲜食机：权限、2.4G Wi‑Fi 与配网")
    doc.add_heading("6.1 绑定前必须满足", level=2)
    add_bullet(doc, "手机蓝牙已开启。")
    add_bullet(doc, "手机已连接可用的 2.4G Wi‑Fi；鲜食机不使用 5G Wi‑Fi 配网。")
    add_bullet(doc, "手机靠近鲜食机和路由器，建议 1–3 米范围内。")
    add_bullet(doc, "APP 已获得定位、附近设备/蓝牙等系统权限。")
    add_bullet(doc, "不要先在手机系统蓝牙设置中手动配对鲜食机，应从 APP 内完成绑定。")
    add_callout(doc, "2.4G 是 Wi‑Fi 频段，不是移动网络", "路由器可能同时广播 2.4G 与 5G。请选择 2.4G 网络；若两个频段同名且配网失败，可在路由器后台暂时分开名称或暂时关闭 5G。", "warn")

    doc.add_heading("6.2 打开手机权限", level=2)
    add_step(doc, 1, "进入烹饪中心", "点击底部“烹饪”，再点击右上角“+”添加鲜食机。")
    add_step(doc, 2, "允许位置信息", "系统弹出定位权限时，选择“仅使用期间允许”或“仅本次使用”。拒绝后可能无法读取 Wi‑Fi 名称或扫描设备。")
    add_figure(doc, "sdk-location-permission.jpg", "图 6-1　Android 定位权限弹窗")
    add_step(doc, 3, "允许查找附近设备", "系统询问是否允许查找、连接附近设备时点击“允许”。")
    add_figure(doc, "sdk-nearby-device-permission.jpg", "图 6-2　Android 附近设备/蓝牙权限弹窗")
    add_callout(doc, "曾经点过“禁止”怎么办", "进入手机“设置 → 应用 → Pet Chef → 权限”，打开位置信息与附近设备；返回 APP 后重新扫描。", "info")

    doc.add_heading("6.3 让鲜食机进入配网模式", level=2)
    add_step(doc, 1, "打开电源", "确认鲜食机通电并处于待机状态。")
    add_step(doc, 2, "长按组合键", "同时长按温度键和功率键约 5 秒。")
    add_step(doc, 3, "确认指示灯闪烁", "看到配网指示灯闪烁后，在 APP 中开始扫描。")
    add_figure(doc, "sdk-device-scan.jpg", "图 6-3　APP 扫描附近鲜食机；点击“选择”")

    doc.add_heading("6.4 选择 2.4G Wi‑Fi 并绑定", level=2)
    add_step(doc, 1, "检查 Wi‑Fi 名称", "APP 会读取手机当前 Wi‑Fi。若显示“未获取到 Wi‑Fi 名称”，优先检查定位权限和手机是否已连接 Wi‑Fi。")
    add_figure(doc, "sdk-wifi-select.jpg", "图 6-4　选择 2.4G Wi‑Fi；未读取名称时先检查定位权限")
    add_step(doc, 2, "输入 Wi‑Fi 密码", "核对网络名称后点击“下一步”，输入该 2.4G Wi‑Fi 的密码。")
    add_figure(doc, "sdk-wifi-password-empty.jpg", "图 6-5　Wi‑Fi 密码输入页")
    add_step(doc, 3, "检查密码", "可短暂点击“显示”核对大小写和字符，确认后再隐藏；不要在公共场所展示密码。")
    add_figure(doc, "sdk-wifi-password-filled.jpg", "图 6-6　密码填写完成后，“一键绑定”按钮变为可用")
    add_step(doc, 4, "点击一键绑定", "保持手机、设备和路由器位置不变，不要退出 APP 或切换网络。")
    add_figure(doc, "sdk-binding-progress.jpg", "图 6-7　绑定依次完成：连接设备、发送 Wi‑Fi、连接云端、绑定账号")
    add_step(doc, 5, "确认绑定成功", "设备出现在“我的鲜食机”列表并显示在线/待机，即可进入设备页面。")

    doc.add_heading("6.5 常见配网失败处理", level=2)
    add_bullet(doc, "找不到设备：重新长按温度键+功率键 5 秒，确认指示灯闪烁。")
    add_bullet(doc, "读取不到 Wi‑Fi：打开定位权限，连接 2.4G Wi‑Fi 后重试。")
    add_bullet(doc, "密码错误：删除已填密码重新输入，区分大小写。")
    add_bullet(doc, "卡在连接云端：检查路由器能否上网，关闭 VPN/代理后重试。")
    add_bullet(doc, "双频路由器：将 2.4G 与 5G 临时改为不同名称，再选择 2.4G。")
    add_bullet(doc, "仍失败：重启鲜食机、路由器和 APP，靠近设备后重新配网。")

    # Chapter 7
    chapter(doc, "第 7 章　鲜食评估：粘贴识别或手动输入食谱")
    doc.add_heading("7.1 进入鲜食验证", level=2)
    add_step(doc, 1, "打开入口", "点击首页“AI 鲜食评估”，或在食谱相关页面进入“鲜食验证”。")
    add_step(doc, 2, "选择宠物", "从下拉列表选择本次食谱要喂给的宠物。多宠物家庭必须在每次评估前核对。")
    add_figure(doc, "app-fresh-check-input.png", "图 7-1　鲜食验证：先选择宠物，再输入或粘贴食谱")

    doc.add_heading("7.2 从小红书或其他来源复制食谱", level=2)
    add_step(doc, 1, "复制有效内容", "复制食材名称和明确克数，例如“牛肉 180 克、南瓜 80 克、西兰花 40 克”。")
    add_step(doc, 2, "粘贴到智能识别编辑区", "长按编辑框选择“粘贴”。可删除作者介绍、营销文案和与配方无关的文字。")
    add_step(doc, 3, "点击“智能识别”", "APP 把明确的食材+克数转换成待验证食材行。")
    add_step(doc, 4, "逐项核对", "确认名称与克数没有识别错误。AI 不应替用户猜测缺失克数；缺失项需手动补充。")
    add_callout(doc, "复制食谱的风险", "社交平台食谱不等于完整日粮。即使原文写有“营养均衡”，也应检查过敏、能量、长期适用性、维矿补充和烹饪方法。", "warn")

    doc.add_heading("7.3 手动逐项输入", level=2)
    add_step(doc, 1, "新增食材", "在“待验证食谱”区域点击“+ 新增食材”。")
    add_step(doc, 2, "填写名称", "每行只填一种食材，例如“牛肉”，不要把多种食材写在同一行。")
    add_step(doc, 3, "填写克数", "输入大于 0 的数值，单位为克。")
    add_step(doc, 4, "删除错误行", "点击对应行左侧“删除”。")
    add_step(doc, 5, "核对总重量", "检查总重量是否与原食谱或实际称量一致。")
    add_step(doc, 6, "开始验证", "至少有一个有效食材并已选择宠物后，点击“验证食材”。")

    doc.add_heading("7.4 如何看评估结果", level=2)
    add_bullet(doc, "优先处理红色“必须处理”风险，再看黄色/提示级调整。")
    add_bullet(doc, "查看每日能量、建议餐次、每餐克数和食谱能量密度。")
    add_bullet(doc, "查看宠物适配评分，了解年龄、体型、体重、活动、健康与过敏等维度。")
    add_bullet(doc, "查看宏量营养结构和数据覆盖说明；估算值不等于实验室检测值。")
    add_bullet(doc, "若提示需要 B 营养包，按建议选择兼容的营养包并确认剂量。")
    add_bullet(doc, "若显示“需要专业确认”，不要作为长期主食直接使用。")
    add_figure(doc, "app-fresh-check-result-actual-part-1.jpg", "图 7-2a　鲜食验证结果（上）：综合评分、风险提示与必须确认事项", width=2.8)
    add_figure(doc, "app-fresh-check-result-actual-part-2.jpg", "图 7-2b　鲜食验证结果：每日营养需求估算", width=3.0)
    add_figure(doc, "app-fresh-check-result-actual-part-3.jpg", "图 7-2c　鲜食验证结果：宠物适配性明细", width=3.0)
    add_figure(doc, "app-fresh-check-result-actual-part-4.jpg", "图 7-2d　鲜食验证结果（下）：宏量结构、长期适宜性与验证结论", width=3.0)

    # Chapter 8
    chapter(doc, "第 8 章　查看 AI 食谱推荐，理解宠物与食谱的关系")
    doc.add_heading("8.1 宠物是分析对象，食谱是被评估方案", level=2)
    add_bullet(doc, "宠物档案提供年龄、体型、体重、BCS、活动、目标、过敏和健康限制。")
    add_bullet(doc, "AI 根据这些信息生成营养需求，再对候选食谱进行排序和提示。")
    add_bullet(doc, "同一食谱面对不同宠物，建议克数、适配度、风险提示和营养包组合可能不同。")
    add_bullet(doc, "更新宠物档案后，应重新运行 AI 分析，不要继续沿用旧截图或旧份量。")
    add_figure(doc, "app-ai-recipe-recommendation.png", "图 8-1　推荐页面把宠物需求与 A+B+C 食谱组合关联起来")

    doc.add_heading("8.2 查看和分享 AI 食谱信息", level=2)
    add_step(doc, 1, "确认宠物名称/品种", "分享前先说明这份建议是为哪只宠物、什么阶段生成。")
    add_step(doc, 2, "查看食谱详情", "打开 A 包详情，记录食材、比例、推荐理由和注意事项。")
    add_step(doc, 3, "说明营养包", "分享时同时说明 B 包为必选补充、C 包为按需支持，避免他人只抄基础食材。")
    add_step(doc, 4, "保留风险提示", "不要裁掉过敏、疾病、能量和专业确认提示。")
    add_step(doc, 5, "分享截图或文字", "使用手机系统截图/分享功能发送给家人或专业人士。APP 当前页面若没有专门“分享”按钮，可使用系统截图分享。")
    add_callout(doc, "不要直接照搬", "别人宠物的 AI 食谱分享只可作为参考。使用前必须为自己的宠物建立档案并重新评估。", "warn")

    # Chapter 9
    chapter(doc, "第 9 章　一键烹饪与喂食反馈")
    doc.add_heading("9.1 开始烹饪前", level=2)
    add_bullet(doc, "鲜食机已绑定、在线并处于待机。")
    add_bullet(doc, "已选择正确宠物和食谱。")
    add_bullet(doc, "食材按 APP 克数称量，清水按提示加入。")
    add_bullet(doc, "锅体、刀组和上盖已正确安装；未超过最大容量。")
    add_bullet(doc, "食材已按要求去核、去骨和切块，避免硬骨、金属或其他异物。")
    add_figure(doc, "app-one-click-cooking.png", "图 9-1　确认食材与份数后点击“一键烹饪”")

    doc.add_heading("9.2 一键烹饪", level=2)
    add_step(doc, 1, "点击“一键烹饪”", "APP 将食谱、份量和烹饪参数带入烹饪中心。")
    add_step(doc, 2, "完成安全确认", "逐项勾选食材、锅体/刀组、上盖等安全检查；上盖未闭合时不能启动。")
    add_step(doc, 3, "确认开始", "点击“开始烹饪”，保持设备通电和联网。")
    add_step(doc, 4, "观察状态", "APP 显示温度、功率、剩余时间和进度。设备离线或状态长时间不更新时，不要重复连续点击启动。")
    add_step(doc, 5, "暂停/继续", "需要暂停时使用 APP 对应按钮；暂停后检查安全，再继续。")
    add_step(doc, 6, "停止", "若必须中止，按页面提示长按停止/重置，避免误触。")
    add_step(doc, 7, "完成与降温", "完成后等待温度下降，断开危险动作，再打开上盖取出食物。")
    add_callout(doc, "远程状态提示", "APP 显示“已发送”不代表设备一定完成动作。以设备实际状态、声音和安全指示为准；异常时先确保现场安全。", "warn")

    doc.add_heading("9.3 喂食与反馈", level=2)
    add_step(doc, 1, "冷却并分装", "食物降至适宜温度后称量本餐份量；其余按食品安全要求冷藏/冷冻。")
    add_step(doc, 2, "完成喂食", "记录实际喂食量，观察宠物是否吃完。")
    add_step(doc, 3, "打开使用记录", "进入“烹饪 → 使用记录”，找到刚完成的烹饪记录。")
    add_step(doc, 4, "点击“喂食反馈”", "选择适口性（宠物接受程度）和便便状态。")
    add_step(doc, 5, "确认提交", "两项都选择后点击“确认”。已提交的记录会显示反馈状态，避免重复提交。")
    add_step(doc, 6, "持续观察", "连续记录 3–7 天；如出现呕吐、持续腹泻、便血、严重瘙痒或精神异常，应停止相关食谱并咨询兽医。")
    add_callout(doc, "反馈的用途", "喂食反馈用于形成宠物的长期饮食记录，帮助后续判断适口性和消化反应；一次反馈不能单独证明食谱长期适合。", "info")

    # Appendix
    chapter(doc, "附录　常见问题与故障排查")
    doc.add_heading("A. 注册与登录", level=2)
    add_bullet(doc, "提示手机号格式错误：确认是 1 开头的 11 位大陆手机号，未含空格或“+86”。")
    add_bullet(doc, "提示验证码格式错误：必须输入短信中的 6 位纯数字。")
    add_bullet(doc, "提示验证码失效：验证码已超过 5 分钟、已成功使用，或连续错误达到 5 次，请重新获取。")
    add_bullet(doc, "用户名不可用：换用 1–18 位中文、英文或数字组合，不能全为数字。")

    doc.add_heading("B. 宠物档案", level=2)
    add_bullet(doc, "找不到品种：缩短关键词，尝试中文品种名；仍无结果时使用其他/自定义。")
    add_bullet(doc, "照片未显示：检查相册权限、图片格式和网络；重新选择一张较小的清晰照片。")
    add_bullet(doc, "AI 无法分析：确认档案已成功保存，宠物卡片存在且网络可用。")

    doc.add_heading("C. 食谱与鲜食评估", level=2)
    add_bullet(doc, "识别不到食材：确保文本包含明确的“食材名 + 克数”。")
    add_bullet(doc, "识别克数错误：手动修改对应行，不要直接提交。")
    add_bullet(doc, "食谱含过敏原：返回更换食谱，或咨询专业人士后调整，不要忽略警示。")

    doc.add_heading("D. 设备绑定与烹饪", level=2)
    add_bullet(doc, "无法扫描：允许定位与附近设备权限，打开蓝牙，让机器重新进入配网模式。")
    add_bullet(doc, "Wi‑Fi 名称为空：连接 Wi‑Fi 并打开定位权限。")
    add_bullet(doc, "2.4G/5G 不确定：查看路由器名称或管理页面；必要时给两个频段设置不同名称。")
    add_bullet(doc, "绑定中断：不要连续重复绑定，先确认机器是否已经出现在“我的鲜食机”。")
    add_bullet(doc, "烹饪状态异常：先确保现场安全，再检查设备联网、电源、上盖与锅体状态。")

    doc.add_heading("E. 建议的日常复查", level=2)
    add_bullet(doc, "幼犬：生长阶段体重变化快，建议至少每 2–4 周更新体重。")
    add_bullet(doc, "成年犬：建议每月记录体重和 BCS，明显变化时重新分析。")
    add_bullet(doc, "老年犬/慢性病宠物：按兽医建议更频繁更新检查结果和健康记录。")
    add_bullet(doc, "过敏与便便：更换食谱后连续记录适口性、皮肤和便便状态。")

    doc.add_paragraph()
    add_callout(
        doc,
        "文档维护说明",
        "本说明书按 2026 年 7 月 Android APP 界面编写。APP 更新后按钮位置或系统权限文字可能变化，但“先建档—再分析—选食谱—称量—绑定设备—安全烹饪—反馈”的主流程保持一致。",
        "tip",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
